from django.shortcuts import render
from .forms import UserRegistrationForm
from django.contrib import messages
from .models import UserRegistrationModel, PasswordResetOTP
from django.conf import settings
import pandas as pd
import random
from django.core.mail import send_mail
from datetime import timedelta
from django.utils import timezone
from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status
from .serializers import UserRegistrationSerializer


# user registration
def UserRegisterActions(request):

    if request.method == 'POST':

        form = UserRegistrationForm(request.POST)

        if form.is_valid():

            print('Data is Valid')

            form.save()

            messages.success(request, 'You have been successfully registered')

            form = UserRegistrationForm()

            return render(request, 'UserRegistrations.html', {'form': form})

        else:

            print("Invalid form")

            messages.error(request, 'Please correct the errors below')

            # IMPORTANT → return form with errors
            return render(request, 'UserRegistrations.html', {'form': form})

    else:

        form = UserRegistrationForm()

    return render(request, 'UserRegistrations.html', {'form': form})

#user login check
def UserLoginCheck(request):
    if request.method == "POST":
        loginid=request.POST.get("loginid")
        password=request.POST.get("pswd")
        print(loginid)
        print(password)
        try:
            check=UserRegistrationModel.objects.get(loginid=loginid,password=password)
            status=check.status
            if status=="activated":
                request.session['id']=check.id
                request.session['loginid']=check.loginid
                request.session['password']=check.password
                request.session['email']=check.email
                return render(request,'users/UserHome.html',{})
            else:
                messages.success(request,"your account not activated")
            return render(request,"UserLogin.html")
        except Exception as e:
            print('=======>',e)
        messages.success(request,'invalid details')
    return render(request,'UserLogin.html',{})


# --- API VIEWS FOR MOBILE APP ---

@api_view(['POST'])
def RegisterAPI(request):
    print(f"DEBUG: Registering user with data: {request.data}")
    serializer = UserRegistrationSerializer(data=request.data)
    if serializer.is_valid():
        serializer.save(status='activated') 
        return Response({
            "message": "User registered successfully",
            "user": serializer.data
        }, status=status.HTTP_201_CREATED)
    print(f"DEBUG: Registration failed: {serializer.errors}")
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

@api_view(['POST'])
def LoginAPI(request):
    loginid = str(request.data.get('loginid', '')).strip()
    password = str(request.data.get('password', '')).strip()
    print(f"DEBUG: Login attempt for loginid: '{loginid}', password: '{password}'")
    
    try:
        user = UserRegistrationModel.objects.get(loginid=loginid, password=password)
        if user.status == "activated":
            print(f"DEBUG: User login SUCCESSFUL for {loginid}")
            return Response({
                "message": "Login successful",
                "user": {
                    "id": user.id,
                    "loginid": user.loginid,
                    "name": user.name,
                    "email": user.email
                }
            }, status=status.HTTP_200_OK)
        else:
            print(f"DEBUG: User found but status is {user.status}")
            return Response({"message": "Account not activated"}, status=status.HTTP_403_FORBIDDEN)
    except UserRegistrationModel.DoesNotExist:
        print(f"DEBUG: User not found or password incorrect for '{loginid}'")
        return Response({"message": "Invalid credentials"}, status=status.HTTP_401_UNAUTHORIZED)


    
@api_view(['POST'])
def AdminLoginAPI(request):
    adminid = str(request.data.get('adminid', '')).strip().lower()
    password = str(request.data.get('password', '')).strip()
    print(f"DEBUG: Admin login attempt for adminid: '{adminid}', password: '{password}'")
    
    if adminid == "admin" and password == "admin":
        print("DEBUG: Admin login SUCCESSFUL")
        return Response({
            "message": "Admin login successful",
            "user": {"name": "Administrator", "role": "admin"}
        }, status=status.HTTP_200_OK)
    print(f"DEBUG: Admin credentials incorrect. Expected 'admin'/'admin', got '{adminid}'/'{password}'")
    return Response({"message": "Invalid admin credentials"}, status=status.HTTP_401_UNAUTHORIZED)


# --- NEW MOBILE API ENDPOINTS ---

@api_view(['POST'])
def ResumeUploadAPI(request):
    """
    Mobile API: Accepts job_description (text) and resume (single PDF file).
    """
    print(f"DEBUG: ResumeUploadAPI called. Data keys: {list(request.data.keys())}, Files keys: {list(request.FILES.keys())}")
    job_description = request.data.get('job_description', '').strip()

    # Support both 'resume' (single, mobile) and 'resume_files' (multi, legacy)
    resume_file = request.FILES.get('resume')
    resume_files = request.FILES.getlist('resume_files') if not resume_file else []
    if resume_file:
        resume_files = [resume_file]

    if not job_description:
        return Response({"error": "Job description is required"}, status=status.HTTP_400_BAD_REQUEST)

    if not resume_files:
        return Response({"error": "Please upload a resume PDF file"}, status=status.HTTP_400_BAD_REQUEST)

    upload_dir = os.path.join(settings.MEDIA_ROOT, "uploads")
    os.makedirs(upload_dir, exist_ok=True)
    fs = FileSystemStorage(location=upload_dir)

    processed_resumes = []
    for f in resume_files:
        if not f.name.lower().endswith('.pdf'):
            continue
        filename = fs.save(f.name, f)
        resume_path = os.path.join(upload_dir, filename)
        resume_text = extract_text_from_pdf(resume_path)
        email, phone = extract_entities(resume_text)
        name = os.path.splitext(f.name)[0]
        processed_resumes.append((name, email, phone, resume_text, resume_path))

    if not processed_resumes:
        return Response({"error": "No valid PDF resume found"}, status=status.HTTP_400_BAD_REQUEST)

    tfidf_vectorizer = TfidfVectorizer()
    job_vector = tfidf_vectorizer.fit_transform([job_description])

    results = []
    for name, email, phone, resume_text, resume_path in processed_resumes:
        resume_vector = tfidf_vectorizer.transform([resume_text])
        similarity = cosine_similarity(job_vector, resume_vector)[0][0] * 100
        similarity = round(similarity, 2)

        if similarity > 20 and email != "Not Found":
            RankedResume.objects.update_or_create(
                email=email,
                defaults={
                    'name': name,
                    'phone_number': phone,
                    'similarity_score': similarity,
                    'status': 'present'
                }
            )

        rel_path = os.path.relpath(resume_path, settings.MEDIA_ROOT)
        resume_url = settings.MEDIA_URL + rel_path.replace('\\', '/')

        results.append({
            "name": name,
            "email": email,
            "phone": phone,
            "score": similarity,
            "similarity_score": similarity,   # also expose for mobile frontend
            "url": resume_url
        })

    results.sort(key=lambda x: x['score'], reverse=True)
    top = results[0]
    return Response({
        "results": results,
        "similarity_score": top['similarity_score'],
        "name": top['name'],
        "email": top['email'],
    }, status=status.HTTP_200_OK)

@api_view(['GET'])
def GetRegisteredUsersAPI(request):
    users = UserRegistrationModel.objects.all().values('id', 'name', 'loginid', 'email', 'mobile', 'state', 'status')
    return Response({"users": list(users)}, status=status.HTTP_200_OK)

@api_view(['POST'])
def ManageUserStatusAPI(request):
    user_id = request.data.get('id')
    action = request.data.get('action') # 'activate' or 'delete'
    
    try:
        user = UserRegistrationModel.objects.get(id=user_id)
        if action == 'activate':
            user.status = 'activated'
            user.save()
            return Response({"message": "User activated successfully"}, status=status.HTTP_200_OK)
        elif action == 'delete':
            user.delete()
            return Response({"message": "User deleted successfully"}, status=status.HTTP_200_OK)
        return Response({"message": "Invalid action"}, status=status.HTTP_400_BAD_REQUEST)
    except UserRegistrationModel.DoesNotExist:
        return Response({"message": "User not found"}, status=status.HTTP_404_NOT_FOUND)

@api_view(['GET'])
def GetUserApplicationsAPI(request):
    # For now return all present resumes as applications
    apps = RankedResume.objects.filter(status='present').values('name', 'email', 'similarity_score', 'created_at')
    return Response({"applications": list(apps)}, status=status.HTTP_200_OK)


def UserHome(request):
    return render(request,"users/UserHome.html",{})


def application(request):
    applications = RankedResume.objects.filter(status='present')
    return render(request, 'users/applications.html', {"results": applications})

def past_applications(request):
    applications = RankedResume.objects.filter(status='past')
    return render(request, 'users/past_applications.html', {"results": applications})

def training(request):
    # Mock data for graph
    data = {
        'labels': ['Resumes Uploaded', 'High Score (>75%)', 'Offer Letters Sent'],
        'values': [
            RankedResume.objects.count(), 
            RankedResume.objects.filter(similarity_score__gt=75).count(),
            RankedResume.objects.filter(status='past').count()
        ]
    }
    return render(request, 'users/training.html', {'graph_data': data})


from django.core.mail import EmailMessage
from django.conf import settings
from django.contrib import messages
from django.shortcuts import render
from docx import Document
from io import BytesIO
from django.utils import timezone
def create_offer_letter(candidate_name, candidate_email, candidate_phone, offer_letter_path):
    """
    Function to dynamically create a customized offer letter using the candidate's details.
    """
    # Create a new Word document (DOCX format)
    doc = Document()
    doc.add_paragraph("XYZ COMPANY")
    doc.add_paragraph("123 Business Road, Suite 100")
    doc.add_paragraph("City, State, Zip Code")
    doc.add_paragraph("Phone: (123) 456-7890")
    doc.add_paragraph("Email: hr@xyzcompany.com")
    doc.add_paragraph("Website: www.xyzcompany.com")
    doc.add_paragraph()

    doc.add_paragraph(f"Date: {timezone.now().strftime('%B %d, %Y')}")
    doc.add_paragraph()

    # Candidate Information
    doc.add_paragraph(f"Dear {candidate_name},")
    doc.add_paragraph(f"Email: {candidate_email}")
    doc.add_paragraph(f"Phone: {candidate_phone}")
    doc.add_paragraph()

    # Offer Letter Content
    doc.add_paragraph(
        f"We are pleased to offer you the position at XYZ Company. After careful consideration, we believe your skills "
        f"and experience align well with our organization's goals, and we are excited to have you join our team."
    )

    doc.add_paragraph("Please review the offer carefully and confirm your acceptance by signing and returning the attached acceptance form.")
    doc.add_paragraph("We look forward to your response and are excited about the possibility of working together.")
    doc.add_paragraph()
    doc.add_paragraph("If you have any questions, please feel free to reach out to me directly at hr@xyzcompany.com.")
    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("[Your Name]")
    doc.add_paragraph("[Your Job Title]")
    doc.add_paragraph("XYZ Company")

    doc.add_paragraph("---")
    doc.add_paragraph("Acceptance of Offer:")
    doc.add_paragraph(f"I, {candidate_name}, hereby accept the offer extended by XYZ Company for the position as outlined in the above letter.")
    doc.add_paragraph("I agree to the terms and conditions of employment.")
    doc.add_paragraph("Signature: _________________________")
    doc.add_paragraph("Date: _________________________")
    doc.add_paragraph("Please return this signed acceptance to HR at hr@xyzcompany.com.")

    # Save the offer letter to a temporary file
    offer_letter_dir = os.path.join(settings.MEDIA_ROOT, "offer_letters")

# Create the directory if it doesn't exist
    if not os.path.exists(offer_letter_dir):
        os.makedirs(offer_letter_dir)

    # Define the file path
    offer_letter_path = os.path.join(offer_letter_dir, f"offer_letter_{candidate_email}.docx")

    # Save the offer letter
    doc.save(offer_letter_path)

    return offer_letter_path


def send_offer_letter_to_candidate(request):
    """
    Django view to handle the POST request and send a custom offer letter to the selected candidate.
    """
    if request.method == 'POST':
        # Get candidate's details from the POST request
        candidate_name = request.POST.get('name')
        candidate_email = request.POST.get('email')
        candidate_phone = '8520852085'

        if not candidate_name or not candidate_email or not candidate_phone:
            messages.error(request, "All fields are required.")
            return render(request, 'users/applications.html')

        # Generate the offer letter
        offer_letter_path = create_offer_letter(candidate_name, candidate_email, candidate_phone, None)

        # Send the offer letter
        subject = 'Job Offer: Congratulations on Your Selection'
        message = (
            f"Dear {candidate_name},\n\n"
            "We are pleased to inform you that you have been selected for the position at XYZ Company.\n"
            "Please find attached your official offer letter, including details about your role, salary, and benefits.\n\n"
            "We are excited to have you join our team and look forward to working with you.\n\n"
            "Best regards,\n"
            "XYZ Company\n"
            "Human Resources Team"
        )

        try:
            # Create the email with the attachment (offer letter)
            email_message = EmailMessage(
                subject=subject,
                body=message,
                from_email=settings.DEFAULT_FROM_EMAIL,
                to=[candidate_email]
            )
            email_message.attach_file(offer_letter_path)
            email_message.send(fail_silently=False)

            # Success message
            rr=RankedResume.objects.get(email=candidate_email)
            rr.status = 'past'
            rr.save()
            messages.success(request, f"Offer letter sent to {candidate_email} successfully!")
            applications=RankedResume.objects.filter(status='present')
            return render(request,'users/applications.html',{"results":applications})
        except Exception as e:
            # Error message if something goes wrong
            messages.error(request, f"Failed to send offer letter to {candidate_email}: {str(e)}")

        applications=RankedResume.objects.all()
        return render(request,'users/applications.html',{"results":applications})


    # If the request is GET or invalid, return the form again
    applications=RankedResume.objects.all()
    return render(request,'users/applications.html',{"results":applications})


# views.py
import os
import re
from django.shortcuts import render
from django.core.files.storage import FileSystemStorage
from django.conf import settings
from .models import RankedResume
import PyPDF2
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Extract text from PDFs
def extract_text_from_pdf(pdf_path):
    with open(pdf_path, "rb") as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() or ""  # Handle None case
        return text
import os
import re
import PyPDF2
from django.shortcuts import render
from django.core.files.storage import FileSystemStorage
from django.conf import settings
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from .models import RankedResume


# =====================================
# Extract Text from PDF
# =====================================
def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        with open(pdf_path, "rb") as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
    except:
        pass
    return text


# =====================================
# Extract Email & Phone
# =====================================
def extract_entities(text):
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b'
    phone_pattern = r'\b\d{10}\b'

    emails = re.findall(email_pattern, text)
    phones = re.findall(phone_pattern, text)

    email = emails[0] if emails else "Not Found"
    phone = phones[0] if phones else "Not Found"

    return email, phone


# =====================================
# Resume Ranking + Save to DB
# =====================================
def index(request):

    if request.method == 'POST':

        job_description = request.POST.get('job_description')
        resume_files = request.FILES.getlist('resume_files')

        upload_dir = os.path.join(settings.MEDIA_ROOT, "uploads")
        os.makedirs(upload_dir, exist_ok=True)

        fs = FileSystemStorage(location=upload_dir)

        processed_resumes = []

        # -----------------------------
        # Save & Process Resumes
        # -----------------------------
        for resume_file in resume_files:

            if not resume_file.name.lower().endswith('.pdf'):
                continue

            filename = fs.save(resume_file.name, resume_file)
            resume_path = os.path.join(upload_dir, filename)

            resume_text = extract_text_from_pdf(resume_path)
            email, phone = extract_entities(resume_text)

            name = os.path.splitext(resume_file.name)[0]

            processed_resumes.append(
                (name, email, phone, resume_text, resume_path)
            )

        # -----------------------------
        # TF-IDF + Cosine Similarity
        # -----------------------------
        tfidf_vectorizer = TfidfVectorizer()
        job_vector = tfidf_vectorizer.fit_transform([job_description])

        ranked_resumes = []

        for name, email, phone, resume_text, resume_path in processed_resumes:

            resume_vector = tfidf_vectorizer.transform([resume_text])

            similarity = cosine_similarity(
                job_vector,
                resume_vector
            )[0][0] * 100

            similarity = round(similarity, 2)

            ranked_resumes.append(
                (name, email, phone, similarity, resume_path)
            )

            # -----------------------------
            # SAVE TO DATABASE (ACTIVE)
            # -----------------------------
            if similarity > 20 and email != "Not Found":

                RankedResume.objects.update_or_create(
                    email=email,
                    defaults={
                        'name': name,
                        'similarity_score': similarity,
                        'status': 'present'
                    }
                )

        # -----------------------------
        # Sort by Highest Similarity
        # -----------------------------
        ranked_resumes.sort(key=lambda x: x[3], reverse=True)

        # -----------------------------
        # Prepare Results for Template
        # -----------------------------
        results = []

        for name, email, phone, similarity, resume_path in ranked_resumes:

            rel_path = os.path.relpath(resume_path, settings.MEDIA_ROOT)
            resume_url = settings.MEDIA_URL + rel_path.replace('\\', '/')

            results.append(
                (name, email, phone, similarity, resume_url)
            )

        return render(
            request,
            'users/prediction_results.html',
            {'results': results}
        )

    return render(request, 'users/upload_resumes.html')


def ForgotPassword(request):
    return render(request, 'ForgotPassword.html')


def SendOTP(request):
    if request.method == "POST":
        email = request.POST.get("email")
        try:
            user = UserRegistrationModel.objects.get(email=email)
            otp = str(random.randint(100000, 999999))
            PasswordResetOTP.objects.filter(email=email, is_used=False).delete()
            PasswordResetOTP.objects.create(email=email, otp=otp)

            subject = 'Password Reset OTP'
            message = f'Your OTP for password reset is: {otp}\n\nThis OTP is valid for 10 minutes.'
            send_mail(subject, message, settings.EMAIL_HOST_USER, [email], fail_silently=False)

            request.session['reset_email'] = email
            messages.success(request, 'OTP sent to your email')
            return render(request, 'VerifyOTP.html', {'email': email})
        except UserRegistrationModel.DoesNotExist:
            messages.error(request, 'Email not registered')
            return render(request, 'ForgotPassword.html')
    return render(request, 'ForgotPassword.html')


def VerifyOTP(request):
    if request.method == "POST":
        email = request.session.get('reset_email')
        otp = request.POST.get("otp")

        try:
            otp_record = PasswordResetOTP.objects.get(email=email, otp=otp, is_used=False)
            time_diff = timezone.now() - otp_record.created_at

            if time_diff > timedelta(minutes=10):
                messages.error(request, 'OTP expired. Please request a new one.')
                return render(request, 'VerifyOTP.html', {'email': email})

            request.session['otp_verified'] = True
            return render(request, 'ResetPassword.html', {'email': email})
        except PasswordResetOTP.DoesNotExist:
            messages.error(request, 'Invalid OTP')
            return render(request, 'VerifyOTP.html', {'email': email})

    email = request.session.get('reset_email')
    return render(request, 'VerifyOTP.html', {'email': email})


def ResetPassword(request):
    if request.method == "POST":
        email = request.session.get('reset_email')
        otp_verified = request.session.get('otp_verified')

        if not otp_verified:
            messages.error(request, 'Please verify OTP first')
            return render(request, 'ForgotPassword.html')

        new_password = request.POST.get("new_password")
        confirm_password = request.POST.get("confirm_password")

        if new_password != confirm_password:
            messages.error(request, 'Passwords do not match')
            return render(request, 'ResetPassword.html', {'email': email})

        try:
            user = UserRegistrationModel.objects.get(email=email)
            user.password = new_password
            user.save()

            PasswordResetOTP.objects.filter(email=email, is_used=False).update(is_used=True)

            del request.session['reset_email']
            del request.session['otp_verified']

            messages.success(request, 'Password reset successfully. Please login with new password.')
            return render(request, 'UserLogin.html')
        except UserRegistrationModel.DoesNotExist:
            messages.error(request, 'User not found')
            return render(request, 'ForgotPassword.html')

    email = request.session.get('reset_email')
    return render(request, 'ResetPassword.html', {'email': email})